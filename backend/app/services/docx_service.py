"""
Генерация docx из шаблонов: подстановка плейсхолдеров {{...}} из данных заказа.
Шаблоны в папке templates/ в корне проекта.
Используется простая замена строк (шаблоны с пробелами в плейсхолдерах, напр. «ФИО продавец»).
"""
from datetime import date
from decimal import Decimal, InvalidOperation
from pathlib import Path
from io import BytesIO
from typing import Dict, Optional

from docx import Document

# Папка шаблонов: корень проекта / templates
_BASE = Path(__file__).resolve().parent.parent.parent.parent
TEMPLATES_DIR = _BASE / "templates"

# Маппинг: имя плейсхолдера в шаблоне (без {{ }}) → ключ в form_data
PLACEHOLDER_TO_FIELD = {
    "ФИО": "client_fio",
    "ФИО дов": "trustee_fio",
    "ФИО_подписант": "client_fio",
    "Действие": None,
    "Дата рождения": "client_birth_date",
    "Дата рождения продавец": "seller_birth_date",
    "Дата рождения дов": "trustee_birth_date",
    "Место рождения": "client_birth_place",
    "Паспорт": "client_passport",
    "Паспорт серия": "client_passport_series",
    "Паспорт номер": "client_passport_number",
    "Паспорт кем выдан": "client_passport_issued_by",
    "Паспорт когда выдан": "client_passport_issued_date",
    "Код подразделения": "client_passport_division_code",
    "Паспорт дов": "trustee_passport",
    "Паспорт дов серия": "trustee_passport_series",
    "Паспорт дов номер": "trustee_passport_number",
    "Паспорт дов кем выдан": "trustee_passport_issued_by",
    "Паспорт дов когда выдан": "trustee_passport_issued_date",
    "Код подразделения дов": "trustee_passport_division_code",
    "Адрес": "client_address",
    "Адрес дов": None,
    "Телефон": "client_phone",
    "Номер телефона": "client_phone",
    "Номер телефона дов": None,
    "ФИО продавец": "seller_fio",
    "Паспорт продавец": "seller_passport",
    "Паспорт продавец серия": "seller_passport_series",
    "Паспорт продавец номер": "seller_passport_number",
    "Паспорт продавец кем выдан": "seller_passport_issued_by",
    "Паспорт продавец когда выдан": "seller_passport_issued_date",
    "Код подразделения продавец": "seller_passport_division_code",
    "Адрес продавец": "seller_address",
    "VIN": "vin",
    "Марка, модель": "brand_model",
    "Тип ТС": "vehicle_type",
    "Год выпуска": "year",
    "Двигатель": "engine",
    "№ шасси (рамы)": "chassis",
    "№ кузова": "body",
    "Цвет": "color",
    "СРТС": "srts",
    "СРТС серия": "srts_series",
    "СРТС номер": "srts_number",
    "СРТС кем выдан": "srts_issued_by",
    "СРТС когда выдан": "srts_issued_date",
    "Гос. Номер": "plate_number",
    "ПТС": "pts",
    "ПТС серия": "pts_series",
    "ПТС номер": "pts_number",
    "ПТС кем выдан": "pts_issued_by",
    "ПТС когда выдан": "pts_issued_date",
    "Сумма ДКП": "summa_dkp",
    "Сумма ДКП прописью": None,
    "Дата ДКП": "dkp_date",
    "Номер договора": "dkp_number",
    "ДКП": "dkp_summary",
    "Название": "client_legal_name",
    "ИНН": "client_inn",
    "ОГРН": "client_ogrn",
    "Сумма госпошлины": "state_duty",
    "Подпись": "client_fio",
    "Подпись продавец": "seller_fio",
    "Подпись дов": "trustee_fio",
    "Масса": "mass",
    "Мощность": "power",
    "ОСАГО": None,
    "Текущая_дата": None,
}

_PASSPORT_PLACEHOLDER_PREFIXES = {
    "Паспорт": "client",
    "Паспорт продавец": "seller",
    "Паспорт дов": "trustee",
}

_REPRESENTATIVE_TEMPLATES = frozenset({"zaiavlenie.docx"})
_NUMBER_REPRESENTATIVE_TEMPLATES = frozenset({"number.docx"})

_DOWNLOAD_LABELS = {
    "akt_pp.docx": "Акт приема-передачи",
    "DKP.docx": "ДКП",
    "dkp_dar.docx": "Договор дарения",
    "dkp_pieces.docx": "ДКП запчасти",
    "doverennost.docx": "Доверенность",
    "mreo.docx": "МРЭО",
    "number.docx": "Заявление на номера",
    "prokuratura.docx": "Прокуратура",
    "zaiavlenie.docx": "Заявление",
    "zaiavlenie_na_nomera.docx": "Заявление на номера",
}

_ONES = {
    0: "",
    1: "один",
    2: "два",
    3: "три",
    4: "четыре",
    5: "пять",
    6: "шесть",
    7: "семь",
    8: "восемь",
    9: "девять",
}
_ONES_FEMININE = {**_ONES, 1: "одна", 2: "две"}
_TEENS = {
    10: "десять",
    11: "одиннадцать",
    12: "двенадцать",
    13: "тринадцать",
    14: "четырнадцать",
    15: "пятнадцать",
    16: "шестнадцать",
    17: "семнадцать",
    18: "восемнадцать",
    19: "девятнадцать",
}
_TENS = {
    2: "двадцать",
    3: "тридцать",
    4: "сорок",
    5: "пятьдесят",
    6: "шестьдесят",
    7: "семьдесят",
    8: "восемьдесят",
    9: "девяносто",
}
_HUNDREDS = {
    1: "сто",
    2: "двести",
    3: "триста",
    4: "четыреста",
    5: "пятьсот",
    6: "шестьсот",
    7: "семьсот",
    8: "восемьсот",
    9: "девятьсот",
}


def _number_applicant_field(form_data: dict, placeholder: str) -> Optional[str]:
    if not form_data.get("trustee_fio"):
        return None
    if placeholder == "ФИО":
        return str(form_data.get("trustee_fio") or "").strip()
    if placeholder == "Паспорт":
        return _full_passport(form_data, "trustee")
    if placeholder in {"Адрес", "Телефон", "Номер телефона"}:
        return ""
    return None


def _fio_initials(value: Optional[str]) -> str:
    """Возвращает расшифровку подписи в формате Фамилия И.О."""
    if not value:
        return ""
    parts = [part for part in str(value).split() if part]
    if len(parts) < 2:
        return str(value).strip()
    suffix = ""
    if len(parts) >= 3 and parts[-1].lower().replace("ё", "е") in {"оглы", "кызы"}:
        suffix = " " + parts[-1].lower()
        parts = parts[:-1]
    initials = "".join(f"{part[0]}." for part in parts[1:] if part)
    return f"{parts[0]} {initials}{suffix}".strip()


def _full_passport(form_data: dict, prefix: str) -> str:
    passport = form_data.get(f"{prefix}_passport")
    series = form_data.get(f"{prefix}_passport_series")
    number = form_data.get(f"{prefix}_passport_number")
    if not passport and series and number:
        passport = f"{series} {number}"

    parts = [str(passport).strip()] if passport else []
    issued_by = form_data.get(f"{prefix}_passport_issued_by")
    issued_date = form_data.get(f"{prefix}_passport_issued_date")
    division_code = form_data.get(f"{prefix}_passport_division_code")
    if issued_by:
        parts.append(f"выдан {str(issued_by).strip()}")
    if issued_date:
        parts.append(str(issued_date).strip())
    if division_code:
        parts.append(f"код подразделения {str(division_code).strip()}")
    return ", ".join(part for part in parts if part)


def _full_vehicle_doc(form_data: dict, prefix: str, *, with_label: bool = False) -> str:
    doc_value = form_data.get(prefix)
    series = form_data.get(f"{prefix}_series")
    number = form_data.get(f"{prefix}_number")
    if not doc_value and series and number:
        doc_value = f"{series} {number}"

    parts = [str(doc_value).strip()] if doc_value else []
    issued_by = form_data.get(f"{prefix}_issued_by")
    issued_date = form_data.get(f"{prefix}_issued_date")
    if issued_by:
        parts.append(f"выдан {str(issued_by).strip()}")
    if issued_date:
        parts.append(str(issued_date).strip())
    value = ", ".join(part for part in parts if part)
    if with_label and value:
        label = "СРТС" if prefix == "srts" else prefix.upper()
        return f"{label} {value}"
    return value


def _plural_ru(number: int, forms: tuple[str, str, str]) -> str:
    n = abs(number) % 100
    if 11 <= n <= 14:
        return forms[2]
    last = n % 10
    if last == 1:
        return forms[0]
    if 2 <= last <= 4:
        return forms[1]
    return forms[2]


def _triad_words(number: int, feminine: bool = False) -> list[str]:
    words: list[str] = []
    hundreds = number // 100
    tens_units = number % 100
    tens = tens_units // 10
    units = tens_units % 10
    if hundreds:
        words.append(_HUNDREDS[hundreds])
    if 10 <= tens_units <= 19:
        words.append(_TEENS[tens_units])
    else:
        if tens:
            words.append(_TENS[tens])
        ones = _ONES_FEMININE if feminine else _ONES
        if units:
            words.append(ones[units])
    return words


def _integer_to_words_ru(number: int) -> str:
    if number == 0:
        return "ноль"
    groups = [
        (10**9, ("миллиард", "миллиарда", "миллиардов"), False),
        (10**6, ("миллион", "миллиона", "миллионов"), False),
        (10**3, ("тысяча", "тысячи", "тысяч"), True),
        (1, ("", "", ""), False),
    ]
    words: list[str] = []
    for divisor, forms, feminine in groups:
        triad = number // divisor
        number %= divisor
        if not triad:
            continue
        words.extend(_triad_words(triad, feminine=feminine))
        if divisor > 1:
            words.append(_plural_ru(triad, forms))
    return " ".join(words)


def _money_words_ru(value: object) -> str:
    try:
        amount = Decimal(str(value or 0)).quantize(Decimal("0.01"))
    except (InvalidOperation, ValueError):
        return ""
    if amount < 0:
        return ""
    rubles = int(amount)
    kopecks = int((amount - Decimal(rubles)) * 100)
    ruble_word = _plural_ru(rubles, ("рубль", "рубля", "рублей"))
    kopeck_word = _plural_ru(kopecks, ("копейка", "копейки", "копеек"))
    return f"{_integer_to_words_ru(rubles)} {ruble_word} {kopecks:02d} {kopeck_word}"


def _signature_fio(form_data: dict, template_name: Optional[str]) -> str:
    if template_name in _REPRESENTATIVE_TEMPLATES | _NUMBER_REPRESENTATIVE_TEMPLATES and form_data.get("trustee_fio"):
        return _fio_initials(form_data.get("trustee_fio"))
    return _fio_initials(form_data.get("client_fio"))


def _signer_full_fio(form_data: dict, template_name: Optional[str]) -> str:
    if template_name in _REPRESENTATIVE_TEMPLATES | _NUMBER_REPRESENTATIVE_TEMPLATES and form_data.get("trustee_fio"):
        return str(form_data.get("trustee_fio") or "").strip()
    return str(form_data.get("client_fio") or "").strip()


def _dkp_statement_value(form_data: dict) -> str:
    summary = str(form_data.get("dkp_summary") or "").strip()
    if summary:
        return summary if summary.upper().startswith("ДКП") else f"ДКП, {summary}"
    parts = ["ДКП"]
    if form_data.get("dkp_date"):
        parts.append(str(form_data["dkp_date"]))
    try:
        summa_dkp = Decimal(str(form_data.get("summa_dkp") or 0))
    except (InvalidOperation, ValueError):
        summa_dkp = Decimal("0")
    if summa_dkp != 0:
        parts.append(str(form_data["summa_dkp"]))
    if form_data.get("dkp_number"):
        parts.append("№ " + str(form_data["dkp_number"]))
    return ", ".join(parts)


def document_download_filename(template_name: str, form_data: Optional[dict]) -> str:
    form_data = form_data or {}
    label = _DOWNLOAD_LABELS.get(template_name, template_name.rsplit(".", 1)[0])
    fio = form_data.get("client_fio") or form_data.get("client_legal_name") or ""
    initials = _fio_initials(str(fio)) if fio else ""
    suffix = f" - {initials}" if initials else ""
    return f"{label}{suffix}.docx"


def _form_data_to_replace_map(
    form_data: Optional[dict],
    doc_date: Optional[date] = None,
    template_name: Optional[str] = None,
) -> Dict[str, str]:
    """Словарь подстановки: «{{ ключ }}» → значение."""
    if not form_data:
        form_data = {}
    doc_date = doc_date or date.today()
    result = {}
    for placeholder, field_key in PLACEHOLDER_TO_FIELD.items():
        value = form_data.get(field_key) if field_key else None
        if placeholder in _PASSPORT_PLACEHOLDER_PREFIXES:
            value = _full_passport(form_data, _PASSPORT_PLACEHOLDER_PREFIXES[placeholder])
        if placeholder == "СРТС":
            value = _full_vehicle_doc(form_data, "srts", with_label=template_name == "zaiavlenie.docx")
        if placeholder == "ПТС":
            value = _full_vehicle_doc(form_data, "pts")
        if template_name in _NUMBER_REPRESENTATIVE_TEMPLATES:
            number_value = _number_applicant_field(form_data, placeholder)
            if number_value is not None:
                value = number_value
        if placeholder == "Сумма ДКП прописью":
            value = _money_words_ru(form_data.get("summa_dkp"))
        if placeholder == "Подпись":
            value = _signature_fio(form_data, template_name)
        if placeholder == "Подпись продавец":
            value = _fio_initials(form_data.get("seller_fio"))
        if placeholder == "Подпись дов":
            value = _fio_initials(form_data.get("trustee_fio"))
        if placeholder == "ФИО_подписант":
            value = _signer_full_fio(form_data, template_name)
        if value is None and placeholder == "Действие":
            value = "НЗ ЗАМЕНИТЬ" if form_data.get("need_plate") else ""
        if value is None and placeholder == "Текущая_дата":
            value = doc_date.strftime("%d.%m.%Y")
        if value is None and placeholder == "Дата ДКП":
            value = doc_date.strftime("%d.%m.%Y")
        if placeholder == "ДКП" and (form_data.get("dkp_date") or form_data.get("summa_dkp") or form_data.get("dkp_number") or form_data.get("dkp_summary")):
            value = _dkp_statement_value(form_data)
        if value is None:
            value = ""
        result[placeholder] = str(value)
    return result


def _replace_in_paragraph(paragraph, replace_map: dict[str, str]) -> None:
    text = paragraph.text
    for key, value in replace_map.items():
        text = text.replace("{{" + key + "}}", value)
        text = text.replace("{{ " + key + " }}", value)
    if text != paragraph.text:
        paragraph.clear()
        paragraph.add_run(text)


def render_docx(template_name: str, form_data: Optional[dict], doc_date: Optional[date] = None) -> bytes:
    """
    Генерирует docx из шаблона (например DKP.docx), подставляя {{ плейсхолдер }} из form_data.
    Возвращает файл как bytes.
    """
    path = TEMPLATES_DIR / template_name
    if not path.is_file():
        raise FileNotFoundError(f"Шаблон не найден: {template_name}")
    doc = Document(str(path))
    replace_map = _form_data_to_replace_map(form_data, doc_date, template_name)
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replace_map)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replace_map)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
