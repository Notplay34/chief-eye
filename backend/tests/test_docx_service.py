from io import BytesIO

from docx import Document

from app.services.docx_service import render_docx


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def _base_form_data() -> dict:
    return {
        "client_fio": "Иванов Иван Иванович",
        "client_birth_date": "11.12.1990",
        "client_birth_place": "г. Волгоград",
        "client_passport": "1814 123456",
        "client_passport_issued_by": "ГУ МВД России по Волгоградской области",
        "client_passport_issued_date": "01.02.2020",
        "client_passport_division_code": "340-001",
        "client_address": "г. Волгоград, ул. Ленина, д. 10",
        "client_phone": "+79991234567",
        "seller_fio": "Петров Пётр Петрович",
        "seller_birth_date": "10.10.1980",
        "seller_passport": "1814 654321",
        "seller_passport_issued_by": "ОУФМС России по Волгоградской области",
        "seller_passport_issued_date": "03.04.2021",
        "seller_passport_division_code": "340-002",
        "seller_address": "г. Михайловка, ул. Советская, д. 5",
        "trustee_fio": "Сидоров Сидор Сидорович",
        "trustee_birth_date": "09.09.1979",
        "trustee_passport": "1814 777777",
        "trustee_passport_issued_by": "ГУ МВД России по Москве",
        "trustee_passport_issued_date": "05.06.2022",
        "trustee_passport_division_code": "770-003",
        "vin": "XTA217230N0000001",
        "brand_model": "Lada Vesta",
        "vehicle_type": "Легковой",
        "year": "2021",
        "engine": "1.6",
        "chassis": "отсутствует",
        "body": "XTA217230N0000001",
        "color": "Белый",
        "srts": "99AA 123456",
        "plate_number": "A001AA34",
        "pts": "78УУ 123456",
        "dkp_date": "23.04.2026",
        "summa_dkp": "850000",
        "state_duty": "500",
    }


def test_dkp_uses_full_seller_passport_without_replacing_seller_with_trustee():
    text = _docx_text(render_docx("DKP.docx", _base_form_data()))

    assert "Петров Пётр Петрович" in text
    assert "Петров Пётр Петрович 10.10.1980" in text
    assert "Иванов Иван Иванович 11.12.1990" in text
    assert "Сидоров Сидор Сидорович" not in text
    assert (
        "1814 654321, выдан ОУФМС России по Волгоградской области, "
        "03.04.2021, код подразделения 340-002"
    ) in text


def test_doverennost_uses_full_trustee_passport():
    text = _docx_text(render_docx("doverennost.docx", _base_form_data()))

    assert "Сидоров Сидор Сидорович" in text
    assert (
        "1814 777777, выдан ГУ МВД России по Москве, "
        "05.06.2022, код подразделения 770-003"
    ) in text


def test_signature_placeholders_use_surname_and_initials():
    form_data = _base_form_data()
    form_data["trustee_fio"] = None

    for template_name in ("number.docx", "mreo.docx", "zaiavlenie.docx"):
        text = _docx_text(render_docx(template_name, form_data))
        assert "Иванов И.И." in text


def test_signature_initials_keep_ogly_kyzy_as_patronymic_suffix():
    form_data = _base_form_data()
    form_data["client_fio"] = "Алиев Руслан Мамед оглы"

    text = _docx_text(render_docx("number.docx", form_data))

    assert "Алиев Р.М. оглы" in text
    assert "Алиев Р.М.О." not in text


def test_zaiavlenie_uses_full_signer_name_and_birth_line():
    text = _docx_text(render_docx("zaiavlenie.docx", _base_form_data()))

    assert "Я, Сидоров Сидор Сидорович" in text
    assert "Я, Сидоров С.С." not in text
    assert "11.12.1990 г. Волгоград" in text


def test_partial_passport_details_do_not_create_extra_separators():
    form_data = {
        "client_fio": "Иванов Иван",
        "client_passport": "1814 123456",
        "client_passport_issued_by": "ГУ МВД России",
        "client_address": "г. Волгоград",
        "client_phone": "+79991234567",
        "brand_model": "Lada Vesta",
        "year": "2021",
    }

    text = _docx_text(render_docx("number.docx", form_data))

    assert "1814 123456, выдан ГУ МВД России" in text
    assert "1814 123456, выдан ГУ МВД России," not in text
    assert ", ," not in text
